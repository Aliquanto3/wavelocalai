"""
Agent Tools - Outils disponibles pour les agents autonomes.

Ce module contient tous les outils que les agents LangGraph et CrewAI peuvent utiliser.
Chaque outil est implémenté selon le pattern :
1. Fonction pure de logique métier (testable)
2. Wrapper LangChain avec décorateur @tool

Nouveaux outils ajoutés :
- Email Sender
- Data Analyzer (CSV/Excel)
- Document Generator (DOCX)
- Chart Generator (PNG)
- Markdown Report Builder
- System Monitor
"""

import datetime
import json
import os
import re
import smtplib
import unicodedata
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from pathlib import Path

import matplotlib
import matplotlib.pyplot as plt
import numexpr
import pandas as pd
import psutil
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from func_timeout import FunctionTimedOut, func_timeout
from langchain_core.tools import tool

# Configuration matplotlib pour éviter les problèmes d'affichage
matplotlib.use("Agg")

# ========================================
# CONFIGURATION
# ========================================

# Répertoire de sortie pour les fichiers générés
OUTPUT_DIR = Path("outputs")
OUTPUT_DIR.mkdir(exist_ok=True)

# Configuration Email (à mettre dans .env en production)
SMTP_SERVER = os.getenv("SMTP_SERVER", "smtp.gmail.com")
SMTP_PORT = int(os.getenv("SMTP_PORT", "587"))
SMTP_USER = os.getenv("SMTP_USER", "")
SMTP_PASSWORD = os.getenv("SMTP_PASSWORD", "")


# ========================================
# HELPERS
# ========================================


def remove_accents(input_str: str) -> str:
    """Normalise une chaîne de caractères."""
    if not isinstance(input_str, str):
        return str(input_str)
    nfkd_form = unicodedata.normalize("NFKD", input_str)
    return "".join([c for c in nfkd_form if not unicodedata.combining(c)]).lower()


def sanitize_filename(filename: str) -> str:
    """Nettoie un nom de fichier pour éviter les problèmes de sécurité."""
    # Retire les caractères dangereux
    clean = re.sub(r'[<>:"/\\|?*]', "", filename)
    # Limite la longueur
    return clean[:100]


# ========================================
# OUTILS EXISTANTS (LEGACY)
# ========================================


def _calculate_safe(expression: str) -> str:
    """
    Logique pure de calcul sécurisé (testable directement).
    """
    # 1. Validation de la longueur
    if len(expression) > 100:
        return "❌ Erreur : Expression trop longue (max 100 caractères)"

    # 2. Nettoyage préventif des espaces multiples
    expression = " ".join(expression.split())

    # 3. Protection contre les expressions vides
    if not expression.strip():
        return "❌ Erreur : Expression vide"

    # 4. Whitelist STRICTE
    if not re.match(r"^[\d\s+\-*/().]+$", expression):
        return "❌ Erreur : Caractères non autorisés. Utilisez uniquement : + - * / ( ) et nombres"

    # 5. Détection d'opérateurs consécutifs
    if re.search(r"[+\-*/]{2,}", expression):
        return "❌ Erreur : Opérateurs consécutifs détectés"

    # 6. Vérification des parenthèses équilibrées
    if expression.count("(") != expression.count(")"):
        return "❌ Erreur : Parenthèses non équilibrées"

    try:
        # 7. Évaluation avec TIMEOUT de 2 secondes
        def _safe_eval():
            return numexpr.evaluate(expression).item()

        result = func_timeout(2, _safe_eval)

        # 8. Validation du résultat
        if not isinstance(result, (int, float)):
            return "❌ Erreur : Résultat invalide"

        # 9. Détection des valeurs spéciales (inf, nan)
        if result == float("inf") or result == float("-inf"):
            return "❌ Erreur : Résultat infini (division par zéro ou overflow)"

        if result != result:  # Test pour NaN
            return "❌ Erreur : Résultat indéfini (NaN)"

        # 10. Formatage du résultat
        if isinstance(result, float):
            if abs(result - round(result)) < 1e-10:
                return str(int(round(result)))
            else:
                return f"{result:.10g}"

        return str(result)

    except FunctionTimedOut:
        return "❌ Erreur : Calcul trop long (timeout 2s). Simplifiez l'expression"

    except ZeroDivisionError:
        return "❌ Erreur : Division par zéro"

    except (ValueError, SyntaxError) as e:
        return f"❌ Erreur de syntaxe : {str(e)}"

    except Exception as e:
        return f"❌ Erreur de calcul : {str(e)}"


def _get_current_time_impl() -> str:
    """Logique pure pour récupérer l'heure."""
    return datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _search_wavestone_impl(query: str) -> str:
    """Logique pure pour la recherche interne."""
    knowledge = {
        "meteo": "Il fait toujours beau dans le Cloud, mais gris à Paris aujourd'hui (12°C).",
        "anael": "Anaël est un consultant IA Senior spécialisé dans le GenAI.",
        "pue": "Le PUE moyen des datacenters Wavestone est de 1.4.",
        "politique": "La politique Green IT impose d'éteindre les GPU le week-end.",
    }

    query = query.lower()
    results = []
    for key, value in knowledge.items():
        if key in query:
            results.append(value)

    if results:
        return "\n".join(results)
    else:
        return "Aucune information trouvée dans la base interne pour cette requête."


# ========================================
# NOUVEAUX OUTILS
# ========================================


def _send_email_impl(to: str, subject: str, body: str) -> str:
    """
    Logique pure d'envoi d'email via SMTP.

    Args:
        to: Adresse email du destinataire
        subject: Sujet de l'email
        body: Corps du message (peut contenir du HTML)

    Returns:
        str: Message de confirmation ou d'erreur
    """
    # Validation des inputs
    if not to or "@" not in to:
        return "❌ Erreur : Adresse email invalide"

    if not subject or len(subject) > 200:
        return "❌ Erreur : Sujet manquant ou trop long (max 200 caractères)"

    if not body or len(body) > 10000:
        return "❌ Erreur : Corps du message manquant ou trop long (max 10000 caractères)"

    # Vérification de la configuration SMTP
    if not SMTP_USER or not SMTP_PASSWORD:
        return "⚠️ Configuration SMTP manquante. Configurez SMTP_USER et SMTP_PASSWORD dans .env"

    try:
        # Création du message
        msg = MIMEMultipart("alternative")
        msg["From"] = SMTP_USER
        msg["To"] = to
        msg["Subject"] = subject

        # Ajout du corps (supporte HTML)
        part = MIMEText(body, "html" if "<" in body else "plain", "utf-8")
        msg.attach(part)

        # Connexion et envoi
        with smtplib.SMTP(SMTP_SERVER, SMTP_PORT) as server:
            server.starttls()
            server.login(SMTP_USER, SMTP_PASSWORD)
            server.send_message(msg)

        return f"✅ Email envoyé avec succès à {to}"

    except smtplib.SMTPAuthenticationError:
        return "❌ Erreur d'authentification SMTP. Vérifiez vos identifiants."

    except smtplib.SMTPException as e:
        return f"❌ Erreur SMTP : {str(e)}"

    except Exception as e:
        return f"❌ Erreur lors de l'envoi de l'email : {str(e)}"


def _analyze_csv_impl(filepath: str, query: str) -> str:
    """
    Analyse un fichier CSV/Excel et répond à une question.

    Args:
        filepath: Chemin vers le fichier CSV/Excel
        query: Question d'analyse (ex: "moyenne de la colonne X")

    Returns:
        str: Résultats de l'analyse en format texte
    """
    # Validation du chemin
    path = Path(filepath)
    if not path.exists():
        return f"❌ Fichier non trouvé : {filepath}"

    if path.suffix.lower() not in [".csv", ".xlsx", ".xls"]:
        return "❌ Format non supporté. Utilisez .csv, .xlsx ou .xls"

    try:
        # Lecture du fichier
        df = pd.read_csv(filepath) if path.suffix.lower() == ".csv" else pd.read_excel(filepath)

        # Informations de base
        info = f"📊 **Analyse de {path.name}**\n\n"
        info += f"- Lignes : {len(df)}\n"
        info += f"- Colonnes : {len(df.columns)}\n"
        info += f"- Colonnes disponibles : {', '.join(df.columns.tolist())}\n\n"

        # Traitement de la requête
        query_lower = query.lower()

        # Cas 1 : Aperçu des données
        if any(word in query_lower for word in ["aperçu", "preview", "affiche", "montre", "show"]):
            info += "**Aperçu (5 premières lignes) :**\n"
            info += df.head().to_string()

        # Cas 2 : Statistiques descriptives
        elif any(word in query_lower for word in ["stats", "statistiques", "describe", "résumé"]):
            info += "**Statistiques descriptives :**\n"
            info += df.describe().to_string()

        # Cas 3 : Moyenne d'une colonne
        elif "moyenne" in query_lower or "mean" in query_lower:
            # Extraction du nom de colonne (heuristique simple)
            for col in df.columns:
                if col.lower() in query_lower:
                    if pd.api.types.is_numeric_dtype(df[col]):
                        mean_val = df[col].mean()
                        info += f"**Moyenne de '{col}' :** {mean_val:.2f}"
                    else:
                        info += f"❌ La colonne '{col}' n'est pas numérique"
                    break
            else:
                info += "⚠️ Aucune colonne spécifique détectée. Voici les moyennes de toutes les colonnes numériques :\n"
                info += df.mean(numeric_only=True).to_string()

        # Cas 4 : Somme
        elif "somme" in query_lower or "sum" in query_lower or "total" in query_lower:
            for col in df.columns:
                if col.lower() in query_lower:
                    if pd.api.types.is_numeric_dtype(df[col]):
                        sum_val = df[col].sum()
                        info += f"**Somme de '{col}' :** {sum_val:.2f}"
                    else:
                        info += f"❌ La colonne '{col}' n'est pas numérique"
                    break
            else:
                info += "⚠️ Aucune colonne spécifique détectée. Voici les sommes :\n"
                info += df.sum(numeric_only=True).to_string()

        # Cas 5 : Comptage
        elif "compte" in query_lower or "count" in query_lower or "nombre" in query_lower:
            for col in df.columns:
                if col.lower() in query_lower:
                    count = df[col].value_counts()
                    info += f"**Comptage de '{col}' :**\n{count.to_string()}"
                    break
            else:
                info += f"**Nombre total de lignes :** {len(df)}"

        # Cas par défaut : Info générale
        else:
            info += "⚠️ Requête non comprise. Reformulez avec : 'aperçu', 'moyenne de X', 'somme de X', 'stats', etc."

        return info

    except pd.errors.EmptyDataError:
        return "❌ Fichier vide"

    except Exception as e:
        return f"❌ Erreur lors de l'analyse : {str(e)}"


def _generate_document_impl(title: str, content: str, filename: str = None) -> str:
    """
    Génère un document Word (.docx) professionnel.

    Args:
        title: Titre du document
        content: Contenu en Markdown (# Titre, ## Sous-titre, paragraphes)
        filename: Nom du fichier (optionnel, généré automatiquement sinon)

    Returns:
        str: Chemin du fichier généré
    """
    try:
        # Création du document
        doc = Document()

        # Style du titre principal
        title_para = doc.add_heading(title, level=0)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Ajout de la date
        date_para = doc.add_paragraph(
            f"Généré le {datetime.datetime.now().strftime('%d/%m/%Y à %H:%M')}"
        )
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()  # Espace

        # Traitement du contenu Markdown
        lines = content.split("\n")

        for line in lines:
            line = line.strip()

            if not line:
                continue

            # Titre niveau 1
            if line.startswith("# "):
                doc.add_heading(line[2:], level=1)

            # Titre niveau 2
            elif line.startswith("## "):
                doc.add_heading(line[3:], level=2)

            # Titre niveau 3
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)

            # Liste à puces
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")

            # Liste numérotée
            elif re.match(r"^\d+\.\s", line):
                doc.add_paragraph(re.sub(r"^\d+\.\s", "", line), style="List Number")

            # Paragraphe normal
            else:
                doc.add_paragraph(line)

        # Génération du nom de fichier
        if not filename:
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"document_{timestamp}.docx"
        else:
            filename = sanitize_filename(filename)
            if not filename.endswith(".docx"):
                filename += ".docx"

        # Sauvegarde
        filepath = OUTPUT_DIR / filename
        doc.save(str(filepath))

        return f"✅ Document créé : {filepath}"

    except Exception as e:
        return f"❌ Erreur lors de la création du document : {str(e)}"


def _generate_chart_impl(
    data_json: str, chart_type: str = "bar", title: str = "Chart", filename: str = None
) -> str:
    """
    Génère un graphique à partir de données JSON.

    Args:
        data_json: Données au format JSON (ex: '{"labels": ["A", "B"], "values": [10, 20]}')
        chart_type: Type de graphique ('bar', 'line', 'pie')
        title: Titre du graphique
        filename: Nom du fichier (optionnel)

    Returns:
        str: Chemin du fichier PNG généré
    """
    try:
        # Parse des données
        data = json.loads(data_json)

        if "labels" not in data or "values" not in data:
            return '❌ Format JSON invalide. Utilisez : {"labels": [...], "values": [...]}'

        labels = data["labels"]
        values = data["values"]

        if len(labels) != len(values):
            return "❌ Le nombre de labels et de valeurs doit être identique"

        # Création du graphique
        plt.figure(figsize=(10, 6))

        if chart_type == "bar":
            plt.bar(labels, values, color="#3498db")
            plt.ylabel("Valeurs")

        elif chart_type == "line":
            plt.plot(labels, values, marker="o", linewidth=2, color="#2ecc71")
            plt.ylabel("Valeurs")
            plt.grid(True, alpha=0.3)

        elif chart_type == "pie":
            plt.pie(values, labels=labels, autopct="%1.1f%%", startangle=90)

        else:
            return (
                f"❌ Type de graphique non supporté : {chart_type}. Utilisez 'bar', 'line' ou 'pie'"
            )

        plt.title(title, fontsize=14, fontweight="bold")
        plt.tight_layout()

        # Génération du nom de fichier
        if not filename:
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"chart_{timestamp}.png"
        else:
            filename = sanitize_filename(filename)
            if not filename.endswith(".png"):
                filename += ".png"

        # Sauvegarde
        filepath = OUTPUT_DIR / filename
        plt.savefig(filepath, dpi=150, bbox_inches="tight")
        plt.close()

        return f"✅ Graphique créé : {filepath}"

    except json.JSONDecodeError:
        return "❌ Format JSON invalide"

    except Exception as e:
        return f"❌ Erreur lors de la création du graphique : {str(e)}"


def _generate_markdown_report_impl(title: str, sections: str, filename: str = None) -> str:
    """
    Génère un rapport Markdown structuré.

    Args:
        title: Titre du rapport
        sections: Contenu au format JSON (ex: '{"Introduction": "...", "Analyse": "..."}')
        filename: Nom du fichier (optionnel)

    Returns:
        str: Chemin du fichier MD généré
    """
    try:
        # Parse des sections
        sections_dict = json.loads(sections) if sections.startswith("{") else {"Contenu": sections}

        # Construction du rapport
        report = f"# {title}\n\n"
        report += f"*Généré le {datetime.datetime.now().strftime('%d/%m/%Y à %H:%M')}*\n\n"
        report += "---\n\n"

        # Ajout des sections
        for section_title, section_content in sections_dict.items():
            report += f"## {section_title}\n\n"
            report += f"{section_content}\n\n"

        # Génération du nom de fichier
        if not filename:
            timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
            filename = f"report_{timestamp}.md"
        else:
            filename = sanitize_filename(filename)
            if not filename.endswith(".md"):
                filename += ".md"

        # Sauvegarde
        filepath = OUTPUT_DIR / filename
        with open(filepath, "w", encoding="utf-8") as f:
            f.write(report)

        return f"✅ Rapport Markdown créé : {filepath}"

    except json.JSONDecodeError:
        return "❌ Format JSON invalide pour les sections"

    except Exception as e:
        return f"❌ Erreur lors de la création du rapport : {str(e)}"


def _system_monitor_impl() -> str:
    """
    Récupère les métriques système actuelles.

    Returns:
        str: Rapport détaillé des ressources système
    """
    try:
        # CPU
        cpu_percent = psutil.cpu_percent(interval=1)
        cpu_count = psutil.cpu_count(logical=False)
        cpu_count_logical = psutil.cpu_count(logical=True)

        # RAM
        mem = psutil.virtual_memory()
        ram_total_gb = mem.total / (1024**3)
        ram_used_gb = mem.used / (1024**3)
        ram_available_gb = mem.available / (1024**3)
        ram_percent = mem.percent

        # Disque
        disk = psutil.disk_usage("/")
        disk_total_gb = disk.total / (1024**3)
        disk_used_gb = disk.used / (1024**3)
        disk_free_gb = disk.free / (1024**3)
        disk_percent = disk.percent

        # Construction du rapport
        report = "📊 **Monitoring Système**\n\n"

        report += "### CPU\n"
        report += f"- Utilisation : {cpu_percent}%\n"
        report += f"- Cœurs physiques : {cpu_count}\n"
        report += f"- Cœurs logiques : {cpu_count_logical}\n\n"

        report += "### RAM\n"
        report += f"- Total : {ram_total_gb:.2f} GB\n"
        report += f"- Utilisée : {ram_used_gb:.2f} GB ({ram_percent}%)\n"
        report += f"- Disponible : {ram_available_gb:.2f} GB\n\n"

        report += "### Disque (/)\n"
        report += f"- Total : {disk_total_gb:.2f} GB\n"
        report += f"- Utilisé : {disk_used_gb:.2f} GB ({disk_percent}%)\n"
        report += f"- Libre : {disk_free_gb:.2f} GB\n\n"

        # Alertes
        if ram_percent > 90:
            report += "⚠️ **ALERTE** : RAM critique (>90%)\n"
        elif ram_percent > 80:
            report += "⚠️ RAM élevée (>80%)\n"

        if disk_percent > 90:
            report += "⚠️ **ALERTE** : Disque critique (>90%)\n"
        elif disk_percent > 80:
            report += "⚠️ Disque élevé (>80%)\n"

        return report

    except Exception as e:
        return f"❌ Erreur lors du monitoring : {str(e)}"


# ========================================
# WRAPPERS LANGCHAIN (Export pour agents)
# ========================================


@tool
def get_current_time():
    """Retourne la date et l'heure actuelle précise au format YYYY-MM-DD HH:MM:SS."""
    return _get_current_time_impl()


@tool
def calculator(expression: str) -> str:
    """
    Effectue un calcul mathématique sécurisé avec timeout et validation stricte.

    Limitations de sécurité :
    - Longueur max : 100 caractères
    - Timeout : 2 secondes
    - Opérateurs autorisés : + - * / ( ) . (espaces et chiffres)

    Exemples valides :
    - "2 + 2"
    - "3.14 * (12/4)"
    - "100 / 3"

    Args:
        expression: Expression mathématique à calculer

    Returns:
        str: Résultat du calcul ou message d'erreur
    """
    return _calculate_safe(expression)


@tool
def search_wavestone_internal(query: str) -> str:
    """
    Simule un moteur de recherche interne à l'entreprise Wavestone.
    Utilise cet outil pour chercher des informations sur les employés, les projets ou les politiques RH.

    Args:
        query: Terme de recherche

    Returns:
        str: Résultats de la recherche
    """
    return _search_wavestone_impl(query)


@tool
def send_email(to: str, subject: str, body: str) -> str:
    """
    Envoie un email via SMTP.

    IMPORTANT : Requiert la configuration SMTP dans le fichier .env :
    - SMTP_SERVER
    - SMTP_PORT
    - SMTP_USER
    - SMTP_PASSWORD

    Args:
        to: Adresse email du destinataire
        subject: Sujet de l'email (max 200 caractères)
        body: Corps du message (supporte HTML, max 10000 caractères)

    Returns:
        str: Message de confirmation ou d'erreur

    Exemple:
        send_email("user@example.com", "Rapport d'analyse", "Voici les résultats...")
    """
    return _send_email_impl(to, subject, body)


@tool
def analyze_csv(filepath: str, query: str) -> str:
    """
    Analyse un fichier CSV ou Excel et répond à une question spécifique.

    Formats supportés : .csv, .xlsx, .xls

    Types de requêtes supportées :
    - "aperçu" : Affiche les 5 premières lignes
    - "stats" : Statistiques descriptives complètes
    - "moyenne de [colonne]" : Calcule la moyenne d'une colonne
    - "somme de [colonne]" : Calcule la somme d'une colonne
    - "compte [colonne]" : Comptage des valeurs uniques

    Args:
        filepath: Chemin vers le fichier CSV/Excel
        query: Question d'analyse

    Returns:
        str: Résultats de l'analyse

    Exemple:
        analyze_csv("data/sales.csv", "moyenne de revenue")
    """
    return _analyze_csv_impl(filepath, query)


@tool
def generate_document(title: str, content: str, filename: str = None) -> str:
    """
    Génère un document Word (.docx) professionnel.

    Le contenu peut utiliser la syntaxe Markdown :
    - # Titre niveau 1
    - ## Titre niveau 2
    - ### Titre niveau 3
    - - Liste à puces
    - 1. Liste numérotée

    Args:
        title: Titre principal du document
        content: Contenu en Markdown
        filename: Nom du fichier (optionnel, généré automatiquement sinon)

    Returns:
        str: Chemin du fichier créé

    Exemple:
        generate_document("Rapport Q4", "## Introduction\\n\\nVoici les résultats...")
    """
    return _generate_document_impl(title, content, filename)


@tool
def generate_chart(
    data_json: str, chart_type: str = "bar", title: str = "Chart", filename: str = None
) -> str:
    """
    Génère un graphique (PNG) à partir de données structurées.

    Types de graphiques supportés :
    - "bar" : Diagramme en barres
    - "line" : Courbe
    - "pie" : Camembert

    Args:
        data_json: Données au format JSON {"labels": [...], "values": [...]}
        chart_type: Type de graphique ('bar', 'line', 'pie')
        title: Titre du graphique
        filename: Nom du fichier (optionnel)

    Returns:
        str: Chemin du fichier PNG créé

    Exemple:
        generate_chart('{"labels": ["Jan", "Feb", "Mar"], "values": [10, 25, 15]}', "bar", "Ventes Q1")
    """
    return _generate_chart_impl(data_json, chart_type, title, filename)


@tool
def generate_markdown_report(title: str, sections: str, filename: str = None) -> str:
    """
    Génère un rapport structuré au format Markdown.

    Args:
        title: Titre du rapport
        sections: Sections au format JSON {"Section 1": "contenu...", "Section 2": "contenu..."}
        filename: Nom du fichier (optionnel)

    Returns:
        str: Chemin du fichier MD créé

    Exemple:
        generate_markdown_report("Analyse IA", '{"Introduction": "...", "Résultats": "..."}')
    """
    return _generate_markdown_report_impl(title, sections, filename)


@tool
def system_monitor() -> str:
    """
    Récupère les métriques système actuelles (CPU, RAM, Disque).

    Utilise cet outil pour :
    - Vérifier les ressources disponibles avant une tâche intensive
    - Diagnostiquer des problèmes de performance
    - Générer des rapports de monitoring

    Returns:
        str: Rapport détaillé des métriques système avec alertes si nécessaire

    Exemple d'utilisation:
        "Vérifie l'état du système avant de lancer le benchmark"
    """
    return _system_monitor_impl()


# ========================================
# REGISTRE DES OUTILS
# ========================================

# Liste complète des outils disponibles (pour backward compatibility)
AVAILABLE_TOOLS = [
    get_current_time,
    calculator,
    search_wavestone_internal,
    send_email,
    analyze_csv,
    generate_document,
    generate_chart,
    generate_markdown_report,
    system_monitor,
]

# Métadonnées des outils pour l'UI
TOOLS_METADATA = {
    "get_current_time": {
        "name": "🕒 Time",
        "description": "Heure système",
        "category": "system",
        "requires_config": False,
    },
    "calculator": {
        "name": "🧮 Calculator",
        "description": "Calculs mathématiques",
        "category": "computation",
        "requires_config": False,
    },
    "search_wavestone_internal": {
        "name": "🏢 Wavestone Search",
        "description": "Base interne simulée",
        "category": "data",
        "requires_config": False,
    },
    "send_email": {
        "name": "📧 Email Sender",
        "description": "Envoi d'emails",
        "category": "communication",
        "requires_config": True,
        "config_vars": ["SMTP_SERVER", "SMTP_USER", "SMTP_PASSWORD"],
    },
    "analyze_csv": {
        "name": "📊 Data Analyzer",
        "description": "Analyse CSV/Excel",
        "category": "data",
        "requires_config": False,
    },
    "generate_document": {
        "name": "📝 Document Generator",
        "description": "Création de DOCX",
        "category": "output",
        "requires_config": False,
    },
    "generate_chart": {
        "name": "📈 Chart Generator",
        "description": "Graphiques PNG",
        "category": "output",
        "requires_config": False,
    },
    "generate_markdown_report": {
        "name": "📋 Markdown Report",
        "description": "Rapports MD",
        "category": "output",
        "requires_config": False,
    },
    "system_monitor": {
        "name": "💾 System Monitor",
        "description": "Métriques système",
        "category": "system",
        "requires_config": False,
    },
}


def get_tools_by_names(tool_names: list[str]) -> list:
    """
    Récupère une liste d'outils par leurs noms.

    Args:
        tool_names: Liste des noms d'outils (ex: ["calculator", "send_email"])

    Returns:
        list: Liste des outils LangChain correspondants
    """
    tools_dict = {tool.name: tool for tool in AVAILABLE_TOOLS}
    return [tools_dict[name] for name in tool_names if name in tools_dict]
